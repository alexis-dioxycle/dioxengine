"""Deliverable renderers and their inverses.

The structured sections are the source of truth; Word/Excel files are
RENDERED from them (render_docx / render_xlsx - used by the /export routes
and the SharePoint sync). Because we generate the files ourselves, we can
also parse a round-tripped file back into section content (parse_docx /
parse_xlsx) - that is what makes the SharePoint two-way sync possible: a
colleague edits the .xlsx in SharePoint, the sync pulls the rows back into
the document draft. Parsing is schema-driven and best-effort: unknown
sheets/headings/columns are ignored rather than fatal.
"""
import io

from docx import Document as Docx
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Font
from openpyxl.utils import get_column_letter


def doc_number(doc, head) -> str:
    rev = head.version_number if head else 0
    return f"{doc.project.name.upper().replace(' ', '-')}-{doc.node.node_key.upper()}-{rev:03d}"


def _sheet_title(section) -> str:
    return (section.get("title") or section["key"])[:31]


NOTES_SHEET = "Notes"


# ------------------------------------------------------------------- Excel

def render_xlsx(doc, head) -> bytes:
    """One sheet per table section, a 'Notes' sheet for text sections."""
    content = (head.content if head else {}) or {}
    sections = (doc.node.content_schema or {}).get("sections", [])

    wb = Workbook()
    wb.remove(wb.active)
    hdr_font = Font(bold=True, size=9, name="Calibri")
    texts = []
    for s in sections:
        if s.get("type") == "table":
            ws = wb.create_sheet(title=_sheet_title(s))
            cols = s.get("columns", [])
            ws.append([c["label"] for c in cols])
            for cell in ws[1]:
                cell.font = hdr_font
                cell.alignment = Alignment(wrap_text=True, vertical="top")
            for row in content.get(s["key"], []) or []:
                ws.append([row.get(c["key"], "") for c in cols])
            for i, c in enumerate(cols, 1):
                ws.column_dimensions[get_column_letter(i)].width = max(12, min(38, len(c["label"]) + 6))
            ws.freeze_panes = "A2"
        else:
            texts.append((s.get("title") or s["key"], content.get(s["key"], "") or ""))
    if texts:
        ws = wb.create_sheet(title=NOTES_SHEET)
        for title, body in texts:
            ws.append([title])
            ws.cell(row=ws.max_row, column=1).font = hdr_font
            ws.append([body])
            ws.append([])
        ws.column_dimensions["A"].width = 110
        for r in ws.iter_rows():
            r[0].alignment = Alignment(wrap_text=True, vertical="top")
    if not wb.sheetnames:
        wb.create_sheet(title="Empty")

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _cell_str(v) -> str:
    if v is None:
        return ""
    if isinstance(v, float) and v == int(v):
        return str(int(v))
    return str(v)


def parse_xlsx(data: bytes, content_schema: dict) -> dict:
    """Read a (possibly human-edited) workbook produced by render_xlsx back
    into {section_key: value}. Sheets are matched by title; table columns by
    header label; text sections by their title row in the Notes sheet."""
    wb = load_workbook(io.BytesIO(data), data_only=True)
    sections = (content_schema or {}).get("sections", [])
    out = {}

    by_title = {_sheet_title(s): s for s in sections if s.get("type") == "table"}
    for name in wb.sheetnames:
        s = by_title.get(name)
        if not s:
            continue
        ws = wb[name]
        rows_iter = ws.iter_rows(values_only=True)
        header = next(rows_iter, None) or []
        label_to_key = {c["label"]: c["key"] for c in s.get("columns", [])}
        col_keys = [label_to_key.get(_cell_str(h).strip()) for h in header]
        rows = []
        for raw in rows_iter:
            row = {}
            for key, v in zip(col_keys, raw):
                if key:
                    row[key] = _cell_str(v)
            if any(v.strip() for v in row.values()):
                # fill missing columns so validate_rows round-trips
                for c in s.get("columns", []):
                    row.setdefault(c["key"], "")
                rows.append(row)
        out[s["key"]] = rows

    text_sections = {(s.get("title") or s["key"]): s for s in sections if s.get("type") != "table"}
    if text_sections and NOTES_SHEET in wb.sheetnames:
        ws = wb[NOTES_SHEET]
        rows = [(_cell_str(r[0]) if r else "") for r in ws.iter_rows(values_only=True)]
        i = 0
        while i < len(rows):
            s = text_sections.get(rows[i].strip())
            if s:
                body = []
                i += 1
                while i < len(rows) and rows[i].strip() not in text_sections:
                    body.append(rows[i])
                    i += 1
                while body and not body[-1].strip():
                    body.pop()
                out[s["key"]] = "\n".join(body)
            else:
                i += 1
    return out


# -------------------------------------------------------------------- Word

def render_docx(doc, head) -> bytes:
    """Cover block (doc number, revision, dates, author/reviewer), then one
    heading per section with prose or a styled table. Per-type Dioxycle
    templates (PROJECT-CTL-001 style) come next; this generic renderer is
    the default."""
    content = (head.content if head else {}) or {}
    sections = (doc.node.content_schema or {}).get("sections", [])
    rev = head.version_number if head else 0
    doc_no = doc_number(doc, head)

    d = Docx()
    style = d.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    for s in d.sections:
        s.left_margin = s.right_margin = Cm(2)

    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(doc_no)
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    d.add_heading(doc.node.name, level=0)
    if doc.node.description:
        d.add_paragraph(doc.node.description).runs[0].italic = True

    meta = d.add_table(rows=2, cols=4)
    meta.style = "Table Grid"
    meta.alignment = WD_TABLE_ALIGNMENT.LEFT
    cells = [("Project", doc.project.name), ("Revision", str(rev)),
             ("Status", (head.status if head else "empty")),
             ("Date", (head.updated_at.strftime("%Y-%m-%d") if head and head.updated_at else "")),
             ("Author", doc.author_email or f"({doc.node.author_role or 'unassigned'})"),
             ("Reviewer", doc.reviewer_email or f"({doc.node.reviewer_role or 'unassigned'})"),
             ("Approved by", head.reviewed_by or "" if head else ""), ("Doc N°", doc_no)]
    for i, (k, v) in enumerate(cells):
        cell = meta.rows[i // 4].cells[i % 4]
        para = cell.paragraphs[0]
        run = para.add_run(f"{k}\n")
        run.bold = True
        run.font.size = Pt(7.5)
        para.add_run(str(v)).font.size = Pt(9)

    for s in sections:
        d.add_heading(s.get("title") or s["key"], level=2)
        if s.get("type") == "table":
            cols = s.get("columns", [])
            rows = content.get(s["key"], []) or []
            t = d.add_table(rows=1, cols=len(cols) or 1)
            t.style = "Table Grid"
            for i, c in enumerate(cols):
                run = t.rows[0].cells[i].paragraphs[0].add_run(c["label"])
                run.bold = True
                run.font.size = Pt(8)
            for row in rows:
                tr = t.add_row()
                for i, c in enumerate(cols):
                    tr.cells[i].paragraphs[0].add_run(str(row.get(c["key"], "") or "")).font.size = Pt(8.5)
        else:
            for para_text in (content.get(s["key"], "") or "-").split("\n"):
                d.add_paragraph(para_text)

    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def parse_docx(data: bytes, content_schema: dict) -> dict:
    """Read a (possibly human-edited) document produced by render_docx back
    into {section_key: value}. Level-2 headings mark sections; the following
    paragraphs (text) or first table (rows, matched by header labels) belong
    to that section. Anything before the first known heading (cover, meta
    table) is ignored."""
    d = Docx(io.BytesIO(data))
    sections = (content_schema or {}).get("sections", [])
    by_title = {(s.get("title") or s["key"]): s for s in sections}

    # walk body elements in document order
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    body = []
    for child in d.element.body.iterchildren():
        if child.tag.endswith("}p"):
            body.append(Paragraph(child, d))
        elif child.tag.endswith("}tbl"):
            body.append(Table(child, d))

    out = {}
    current = None
    texts = []

    def flush():
        if current and current.get("type") != "table":
            body_text = "\n".join(texts)
            while body_text.endswith("\n"):
                body_text = body_text[:-1]
            out[current["key"]] = "" if body_text.strip() == "-" else body_text

    for el in body:
        if isinstance(el, Paragraph):
            style = (el.style.name or "") if el.style else ""
            if style.startswith("Heading") and el.text.strip() in by_title:
                flush()
                current = by_title[el.text.strip()]
                texts = []
            elif current is not None and current.get("type") != "table" and not style.startswith("Heading"):
                texts.append(el.text)
        elif isinstance(el, Table) and current is not None and current.get("type") == "table" \
                and current["key"] not in out:
            label_to_key = {c["label"]: c["key"] for c in current.get("columns", [])}
            trs = el.rows
            if not trs:
                continue
            col_keys = [label_to_key.get(c.text.strip()) for c in trs[0].cells]
            rows = []
            for tr in trs[1:]:
                row = {}
                for key, cell in zip(col_keys, tr.cells):
                    if key:
                        row[key] = cell.text
                if any(v.strip() for v in row.values()):
                    for c in current.get("columns", []):
                        row.setdefault(c["key"], "")
                    rows.append(row)
            out[current["key"]] = rows
    flush()
    return out
